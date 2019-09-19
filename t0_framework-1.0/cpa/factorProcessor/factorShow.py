# -*- coding: utf-8 -*-
'''
生成因子报告
@Time    : 2019/7/26 16:12
@Author  : Lee
@Email   : lxbadboy@163.com
'''
import os
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import pandas as pd
import numpy as np

from cpa.config import pathSelector
from cpa.config import const


class FactorShow:

    def __init__(self):
        self.factorDataPath = pathSelector.PathSelector.getFactorFilePath()
        self.factorNameList = os.listdir(self.factorDataPath)
        self.reportPath = pathSelector.PathSelector.getFactorReportPath()

    def AllFactorReports(self):
        for i in self.factorNameList:
            self.singleFactorReport(i)


    def singleFactorReport(self, factorName):

        # 创建空白文档,并设置页面
        document = Document()
        sec = document.sections[0]
        sec.left_margin = Inches(0.4)
        sec.right_margin = Inches(0.4)
        sec.top_margin = Inches(0.3)
        sec.bottom_margin = Inches(0.3)

        '''写标题'''
        document.add_heading(r'Factors Report', 0)

        '''因子名称'''
        para = document.add_paragraph('')
        run = para.add_run(r'Factor Name:  ')
        run.font.bold = True
        run.font.size = Pt(13)
        run = para.add_run(r'{0}'.format(factorName))

        '''回测时间'''
        for root, dirs, files in os.walk(os.path.join(self.factorDataPath, factorName, '5min')):
            timeLine = pd.read_hdf(os.path.join(root, files[1])).index  # 通过储存的相应因子的factor h5数据，获取测试的时间周期
        startTime = timeLine[0]
        endTime = timeLine[-1]
        para = document.add_paragraph('')
        run = para.add_run(r'Back-testing Time:  ')
        run.font.bold = True
        run.font.size = Pt(13)
        run = para.add_run(r'{0} -- {1}'.format(startTime, endTime))

        '''写设定的参数'''
        para = document.add_paragraph('')
        run = para.add_run(r'Parameters:      ')
        run.font.bold = True
        run.font.size = Pt(13)
        statisticFileName = factorName + '_Statistic_' + const.DataFrequency.HOUR + '.xls'
        statisticPath = pathSelector.PathSelector.getFactorFilePath(factorName=factorName,  # 因子计算数据的文件路径
                                                                    factorFrequency=const.DataFrequency.HOUR,
                                                                    fileName=statisticFileName)
        df = pd.read_excel(statisticPath, header=0, index_col=0)
        #lag
        run = para.add_run(r'Lag:  ')
        run.font.bold = True
        run = para.add_run(r'{0},   '.format(df.iloc[8,0]))
        #手续费率
        run = para.add_run(r'Transaction-Fee:  ')
        run.font.bold = True
        run = para.add_run(r'{0},    '.format("%.2f%%" % (df.iloc[9,0] * 100)))
        #
        run = para.add_run(r'Asset-Number:  ')
        run.font.bold = True
        run = para.add_run(r'{0}'.format(int(df.iloc[10,0])))

        '''分别写5min，30min，1h三个截面的数据，包括图和excel'''
        for freq in [const.DataFrequency.MINUTE5,const.DataFrequency.MINUTE30,const.DataFrequency.HOUR]:
            #副标题
            para = document.add_paragraph('')
            run = para.add_run('· '+ freq + ' 截面指标:')
            run.font.bold = True
            run.font.size = Pt(12)
            # 加相应的收益图
            figName = factorName + '_Report_' + freq + '.png'
            figPath = pathSelector.PathSelector.getFactorFilePath(factorName=factorName,  # 因子计算数据的文件路径
                                                                  factorFrequency=freq,
                                                                  fileName=figName)
            document.add_picture(figPath, width=Inches(8))
            statisticFileName = factorName + '_Statistic_' + freq + '.xls'
            statisticPath = pathSelector.PathSelector.getFactorFilePath(factorName=factorName,  # 因子计算数据的文件路径
                                                                      factorFrequency=freq,
                                                                      fileName=statisticFileName)
            df=pd.read_excel(statisticPath,header=0,index_col=0)   #读取相应的indicators的表格数据

            #创建表格，设置相关参数
            indicators = df.iloc[:8]
            tableRow = indicators.shape[0] + 1
            tableCol = indicators.shape[1] + 1
            table = document.add_table(rows=tableRow, cols=tableCol, style='Table Grid')   #新建表格
            table.autofit=False
            table.columns[-1].width = Inches(0.8)
            table.columns[0].width = Inches(0.7)
            for i in np.arange(tableRow):
                table.rows[i].height = Inches(0.2)
            digit=3   #设置小数保留位数
            # 写入表格，遍历每一个单元格
            for col in np.arange(tableCol):
                for row in np.arange(tableRow):
                    if (col==0) & (row==0):
                        continue
                    elif (col==0) & (row!=0):
                        para=table.cell(row,col).add_paragraph('')
                        run = para.add_run(str(indicators.index[row-1]))
                        run.font.bold = True
                    elif (col!=0) & (row==0):
                        para = table.cell(row, col).add_paragraph('')
                        run = para.add_run(str(indicators.columns[col-1]))
                        run.font.bold = True
                    else:
                        para = table.cell(row, col).add_paragraph('')
                        if row in [1,3,6,7]:
                            run = para.add_run("%.2f%%" % (indicators.iloc[row-1,col-1] * 100))
                        elif row==2:
                            run = para.add_run(str(np.round(indicators.iloc[row - 1, col - 1], 2)))
                        else:
                            run = para.add_run(str(np.round(indicators.iloc[row-1,col-1],digit)))  #保留小数
                    run.font.size=Pt(10)

            if freq!=const.DataFrequency.HOUR:
                document.add_page_break()    #写完5min和30min截面后，均新添加一页
            else:
                para = document.add_paragraph('')
                para = document.add_paragraph('')
                run = para.add_run(r'【注】：')
                run.font.bold = True
                run.font.size = Pt(12)
                para = document.add_paragraph('')
                run = para.add_run(r'  1. 图中的GroupRet,CumRet,表中的Ret，均为减去交易成本后的最终收益和累计收益')
                para = document.add_paragraph('')
                run = para.add_run(r'  2. 图中的IC和RankIC均为累计的，表中的是平均的')
                para = document.add_paragraph('')
                run = para.add_run(r'  3. 换手率Turn为每组的平均换手率')
                para = document.add_paragraph('')
                run = para.add_run(r'  4. 交易成本Cost为每组的累计交易成本，Num是每一组平均的持仓资产数量')


        #储存word文档
        document.save(os.path.join(self.reportPath, '{0}_report.docx'.format(factorName)))
        print('Report: %s  generated and saved.' % factorName)

if __name__ == '__main__':
    factorShow = FactorShow()
    # factorShow.AllFactorReports()
    factorShow.singleFactorReport('maPanelFactor')
